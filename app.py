# -*- coding: utf-8 -*-
"""
Relatório Técnico – Streamlit
Recursos:
- Formulário completo (metadados, conteúdo, autores, refs, anexos)
- Logo no cabeçalho (PNG/JPG) + largura configurável
- Prévia em Markdown
- Exporta: Markdown (.md), PDF (ReportLab) e DOCX (python-docx)
- Rascunho local (.json), autosave e carregar rascunhos
- Gerador automático de código (ex.: MavipeRTEC001)
- Upload manual e AUTOMÁTICO para Google Drive e GitHub (via st.secrets)
"""

import io, os, base64, json, datetime as dt
from pathlib import Path
from typing import List, Optional, Tuple

import streamlit as st
from pydantic import BaseModel, Field

# ===================== Modelos =====================
class Autor(BaseModel):
    nome: str = ""
    cargo: str = ""
    email: str = ""

class Referencia(BaseModel):
    referencia: str = ""

class Anexo(BaseModel):
    titulo: str = ""
    descricao: str = ""
    link: str = ""

class Relatorio(BaseModel):
    # Metadados
    titulo: str = "Relatório Técnico"
    cliente: str = ""
    projeto: str = ""
    codigo: str = ""  # será preenchido pelo gerador (opcional)
    data: str = dt.date.today().isoformat()
    versao: str = "1.0"
    # Equipe
    autores: List[Autor] = Field(default_factory=lambda: [Autor()])
    aprovador: str = ""
    # Corpo
    resumo_exec: str = ""
    escopo: str = ""
    dados_fontes: str = ""
    metodologia: str = ""
    resultados: str = ""
    discussoes: str = ""
    conclusoes: str = ""
    recomendacoes: str = ""
    # Outros
    referencias: List[Referencia] = Field(default_factory=list)
    anexos: List[Anexo] = Field(default_factory=list)
    observacoes: str = ""

# ===================== Helpers =====================
def to_markdown(r: Relatorio) -> str:
    autores_md = "\n".join([f"- {a.nome} ({a.cargo}) <{a.email}>" for a in r.autores if a.nome.strip()])
    refs_md = "\n".join([f"- {x.referencia}" for x in r.referencias if x.referencia.strip()])
    anexos_md = "\n".join([
        f"- **{a.titulo}** – {a.descricao} {(f'({a.link})' if a.link else '')}" for a in r.anexos if a.titulo.strip()
    ])
    parts = [
        f"# {r.titulo}",
        f"**Cliente:** {r.cliente}  ",
        f"**Projeto:** {r.projeto}  ",
        f"**Código:** {r.codigo}  ",
        f"**Data:** {r.data}  ",
        f"**Versão:** {r.versao}",
        "\n---\n",
        "## Autores",
        autores_md or "(preencher)",
        f"\n**Aprovador:** {r.aprovador or '(preencher)'}\n",
        "\n## Resumo Executivo\n" + (r.resumo_exec or "(preencher)"),
        "\n## Escopo\n" + (r.escopo or "(preencher)"),
        "\n## Dados & Fontes\n" + (r.dados_fontes or "(preencher)"),
        "\n## Metodologia\n" + (r.metodologia or "(preencher)"),
        "\n## Resultados\n" + (r.resultados or "(preencher)"),
        "\n## Discussões\n" + (r.discussoes or "(preencher)"),
        "\n## Conclusões\n" + (r.conclusoes or "(preencher)"),
        "\n## Recomendações\n" + (r.recomendacoes or "(preencher)"),
        "\n## Referências\n" + (refs_md or "(preencher)"),
        "\n## Anexos\n" + (anexos_md or "(preencher)"),
        "\n## Observações\n" + (r.observacoes or ""),
    ]
    return "\n".join(parts)

def get_logo_dims_cm(logo_bytes: bytes, width_cm: float) -> Tuple[float, float]:
    from PIL import Image as PILImage
    img = PILImage.open(io.BytesIO(logo_bytes))
    w, h = img.size
    if not w or not h:
        return width_cm, width_cm * 0.5
    ratio = h / w
    return width_cm, width_cm * ratio

# ===================== Exportadores =====================
def build_pdf(r: Relatorio, logo_bytes: Optional[bytes], logo_width_cm: float) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    def p(text: str, style="BodyText"):
        story.append(Paragraph(text.replace("\n", "<br/>"), styles[style]))
        story.append(Spacer(1, 0.3*cm))

    if logo_bytes:
        w_cm, h_cm = get_logo_dims_cm(logo_bytes, logo_width_cm)
        story.append(RLImage(io.BytesIO(logo_bytes), width=w_cm*cm, height=h_cm*cm))
        story.append(Spacer(1, 0.4*cm))

    p(f"<b>{r.titulo}</b>", "Title")
    p(
        f"Cliente: <b>{r.cliente or '-'}</b><br/>"
        f"Projeto: <b>{r.projeto or '-'}</b><br/>"
        f"Código: <b>{r.codigo or '-'}</b><br/>"
        f"Data: <b>{r.data or '-'}</b><br/>"
        f"Versão: <b>{r.versao or '-'}</b>"
    )

    autores = "<br/>".join([f"- {a.nome} ({a.cargo}) &lt;{a.email}&gt;" for a in r.autores if a.nome.strip()]) or "(preencher)"
    p(f"<b>Autores</b><br/>{autores}")
    p(f"<b>Aprovador</b><br/>{r.aprovador or '(preencher)'}")

    def sec(title, text):
        p(f"<b>{title}</b>")
        p(text or "(preencher)")

    sec("Resumo Executivo", r.resumo_exec)
    sec("Escopo", r.escopo)
    sec("Dados & Fontes", r.dados_fontes)
    sec("Metodologia", r.metodologia)
    sec("Resultados", r.resultados)
    sec("Discussões", r.discussoes)
    sec("Conclusões", r.conclusoes)
    sec("Recomendações", r.recomendacoes)

    refs = "<br/>".join([f"- {x.referencia}" for x in r.referencias if x.referencia.strip()]) or "(preencher)"
    p(f"<b>Referências</b><br/>{refs}")

    anexos = "<br/>".join([
        f"- <b>{a.titulo}</b> – {a.descricao} {(f'({a.link})' if a.link else '')}" for a in r.anexos if a.titulo.strip()
    ]) or "(preencher)"
    p(f"<b>Anexos</b><br/>{anexos}")

    if r.observacoes:
        sec("Observações", r.observacoes)

    doc = doc  # keep reference for linter
    from reportlab.platypus import SimpleDocTemplate as _tmp  # noqa
    # (já foi construído acima)
    story_doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    story_doc.build(story)
    return buf.getvalue()

def build_docx(r: Relatorio, logo_bytes: Optional[bytes], logo_width_cm: float) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if logo_bytes:
        section = doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(logo_bytes), width=Cm(logo_width_cm))

    doc.add_heading(r.titulo or "Relatório Técnico", level=0)

    meta = doc.add_paragraph()
    meta.add_run("Cliente: ").bold = True; meta.add_run(r.cliente or "-")
    meta.add_run("\nProjeto: ").bold = True; meta.add_run(r.projeto or "-")
    meta.add_run("\nCódigo: ").bold = True; meta.add_run(r.codigo or "-")
    meta.add_run("\nData: ").bold = True; meta.add_run(r.data or "-")
    meta.add_run("\nVersão: ").bold = True; meta.add_run(r.versao or "-")

    doc.add_heading("Autores", level=1)
    for a in r.autores:
        if a.nome.strip():
            doc.add_paragraph(f"- {a.nome} ({a.cargo}) <{a.email}>")
    doc.add_paragraph(f"Aprovador: {r.aprovador or '(preencher)'}")

    def sec(title, text):
        doc.add_heading(title, level=1)
        doc.add_paragraph(text or "(preencher)")

    sec("Resumo Executivo", r.resumo_exec)
    sec("Escopo", r.escopo)
    sec("Dados & Fontes", r.dados_fontes)
    sec("Metodologia", r.metodologia)
    sec("Resultados", r.resultados)
    sec("Discussões", r.discussoes)
    sec("Conclusões", r.conclusoes)
    sec("Recomendações", r.recomendacoes)

    doc.add_heading("Referências", level=1)
    if r.referencias:
        for x in r.referencias:
            if x.referencia.strip():
                doc.add_paragraph(f"- {x.referencia}")
    else:
        doc.add_paragraph("(preencher)")

    doc.add_heading("Anexos", level=1)
    if r.anexos:
        for a in r.anexos:
            if a.titulo.strip():
                line = f"- {a.titulo} – {a.descricao}"
                if a.link:
                    line += f" ({a.link})"
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("(preencher)")

    if r.observacoes:
        sec("Observações", r.observacoes)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ===================== Drive/GitHub =====================
def get_drive_service():
    try:
        from google.oauth2.service_account import Credentials
        from googleapiclient.discovery import build
        scopes = ["https://www.googleapis.com/auth/drive.file"]
        sa_info = dict(st.secrets["gcp_service_account"])  # type: ignore
        creds = Credentials.from_service_account_info(sa_info, scopes=scopes)
        service = build("drive", "v3", credentials=creds)
        return service
    except Exception as e:
        st.error(f"Drive não configurado: {e}")
        return None

def drive_upload_bytes(service, folder_id: str, filename: str, data: bytes, mime: str) -> str:
    from googleapiclient.http import MediaIoBaseUpload
    import io as _io
    file_metadata = {"name": filename, "parents": [folder_id]}
    media = MediaIoBaseUpload(_io.BytesIO(data), mimetype=mime, resumable=False)
    f = service.files().create(body=file_metadata, media_body=media, fields="id, webViewLink").execute()
    return f.get("webViewLink", "")

def github_upload_bytes(filename: str, data: bytes, message: str) -> str:
    try:
        gh = st.secrets.get("github", {})
        token = gh.get("token"); repo = gh.get("repo"); branch = gh.get("branch", "main"); base_path = gh.get("base_path", "reports")
        if not token or not repo:
            raise RuntimeError("Token/Repo não configurados em st.secrets['github']")
        url = f"https://api.github.com/repos/{repo}/contents/{base_path}/{filename}"
        payload = {"message": message, "content": base64.b64encode(data).decode("utf-8"), "branch": branch}
        import requests
        r = requests.put(url, json=payload, headers={"Authorization": f"token {token}", "Accept": "application/vnd.github+json"}, timeout=30)
        if r.status_code not in (200, 201):
            raise RuntimeError(f"GitHub API: {r.status_code} – {r.text}")
        j = r.json()
        return j.get("content", {}).get("html_url", "") or j.get("content", {}).get("path", "")
    except Exception as e:
        st.error(f"GitHub upload falhou: {e}")
        return ""

# ===================== Código automático =====================
COUNTER_FILE_DEFAULT = "counter.json"
def next_report_code(prefix: str = "MavipeRTEC", draft_dir: Optional[str] = None) -> str:
    pdir = Path(draft_dir) if draft_dir else Path.cwd() / "drafts"
    pdir.mkdir(parents=True, exist_ok=True)
    cfile = pdir / COUNTER_FILE_DEFAULT
    counter = 0
    if cfile.exists():
        try:
            counter = json.loads(cfile.read_text(encoding="utf-8")).get("counter", 0)
        except Exception:
            pass
    counter += 1
    cfile.write_text(json.dumps({"counter": counter}, ensure_ascii=False, indent=2), encoding="utf-8")
    return f"{prefix}{counter:03d}"

# ===================== UI =====================
st.set_page_config(page_title="Relatório Técnico", page_icon="📝", layout="wide")
st.title("📝 Relatório Técnico – Editor")

with st.sidebar:
    st.header("Ações & Configuração")

    # Estado
    if "rel" not in st.session_state:
        st.session_state.rel = Relatorio()
    rel: Relatorio = st.session_state.rel

    # Código automático
    st.subheader("Código automático")
    code_prefix = st.text_input("Prefixo", value=st.session_state.get("code_prefix","MavipeRTEC"))
    if st.button("🔢 Gerar código"):
        rel.codigo = next_report_code(prefix=code_prefix)
        st.session_state.rel = rel
        st.success(f"Código: {rel.codigo}")

    st.markdown("---")
    st.subheader("Rascunho local")
    default_dir = st.session_state.get("draft_dir", str(Path.cwd() / "drafts"))
    draft_dir = st.text_input("Pasta de rascunhos", value=default_dir)
    st.session_state.draft_dir = draft_dir
    autosave = st.checkbox("Autosave ao atualizar prévia", value=st.session_state.get("autosave", True))
    st.session_state.autosave = autosave

    # Upload automático para cloud
    st.markdown("---")
    st.subheader("Upload automático (cloud)")
    auto_drive = st.checkbox("Google Drive (usar st.secrets)", value=st.session_state.get("auto_drive", False))
    st.session_state.auto_drive = auto_drive
    auto_gh = st.checkbox("GitHub (usar st.secrets)", value=st.session_state.get("auto_gh", False))
    st.session_state.auto_gh = auto_gh

    # Carregar/Salvar local
    try:
        p = Path(draft_dir); p.mkdir(parents=True, exist_ok=True)
        opts = ["(nenhum)"] + [f.name for f in sorted(p.glob("*.json"))]
    except Exception as e:
        opts = ["(nenhum)"]
        st.error(f"Pasta inválida: {e}")
    pick = st.selectbox("Carregar rascunho", opts)
    if pick != "(nenhum)":
        try:
            data = json.loads((Path(draft_dir)/pick).read_text(encoding="utf-8"))
            st.session_state.rel = Relatorio(**data)
            rel = st.session_state.rel
            st.success(f"Carregado: {pick}")
        except Exception as e:
            st.error(f"Falha ao carregar: {e}")

    if st.button("💾 Salvar local agora"):
        try:
            p = Path(draft_dir); p.mkdir(parents=True, exist_ok=True)
            name = f"{(rel.codigo or 'relatorio').replace(' ','_')}.json"
            (p / name).write_text(json.dumps(rel.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")
            st.success(f"Salvo: {p / name}")
        except Exception as e:
            st.error(f"Erro ao salvar: {e}")

    st.markdown("---")
    st.subheader("Logo")
    logo_file = st.file_uploader("Logo (PNG/JPG)", type=["png","jpg","jpeg"])
    logo_width_cm = st.number_input("Largura do logo (cm)", 1.0, 12.0, 3.5, 0.5)
    if "logo_bytes" not in st.session_state:
        st.session_state.logo_bytes = None
    if logo_file:
        st.session_state.logo_bytes = logo_file.read()
        st.success("Logo carregado.")

# -------- Form --------
rel: Relatorio = st.session_state.rel
with st.form("form"):
    st.subheader("Metadados")
    c1,c2,c3,c4 = st.columns([2,2,1,1])
    rel.titulo  = c1.text_input("Título", rel.titulo)
    rel.cliente = c2.text_input("Cliente", rel.cliente)
    rel.projeto = c3.text_input("Projeto", rel.projeto)
    rel.codigo  = c4.text_input("Código", rel.codigo)
    d1,d2 = st.columns(2)
    rel.data   = d1.date_input("Data", dt.date.fromisoformat(rel.data) if rel.data else dt.date.today()).isoformat()
    rel.versao = d2.text_input("Versão", rel.versao)

    st.subheader("Equipe")
    n = st.number_input("Nº de autores", 1, 10, max(1, len(rel.autores)), 1)
    while len(rel.autores) < n: rel.autores.append(Autor())
    while len(rel.autores) > n: rel.autores.pop()
    for i,a in enumerate(rel.autores):
        c1,c2,c3 = st.columns([2,2,2])
        a.nome  = c1.text_input(f"Autor {i+1} – Nome", a.nome)
        a.cargo = c2.text_input(f"Autor {i+1} – Cargo", a.cargo)
        a.email = c3.text_input(f"Autor {i+1} – E-mail", a.email)
    rel.aprovador = st.text_input("Aprovador", rel.aprovador)

    st.subheader("Conteúdo")
    rel.resumo_exec   = st.text_area("Resumo Executivo", rel.resumo_exec)
    rel.escopo        = st.text_area("Escopo", rel.escopo)
    rel.dados_fontes  = st.text_area("Dados & Fontes", rel.dados_fontes)
    rel.metodologia   = st.text_area("Metodologia", rel.metodologia)
    rel.resultados    = st.text_area("Resultados", rel.resultados)
    rel.discussoes    = st.text_area("Discussões", rel.discussoes)
    rel.conclusoes    = st.text_area("Conclusões", rel.conclusoes)
    rel.recomendacoes = st.text_area("Recomendações", rel.recomendacoes)

    st.subheader("Referências & Anexos")
    nr = st.number_input("Nº de referências", 0, 50, len(rel.referencias), 1)
    while len(rel.referencias) < nr: rel.referencias.append(Referencia())
    while len(rel.referencias) > nr: rel.referencias.pop()
    for i,rr in enumerate(rel.referencias):
        rr.referencia = st.text_input(f"Ref. {i+1}", rr.referencia)

    na = st.number_input("Nº de anexos", 0, 30, len(rel.anexos), 1)
    while len(rel.anexos) < na: rel.anexos.append(Anexo())
    while len(rel.anexos) > na: rel.anexos.pop()
    for i,ax in enumerate(rel.anexos):
        c1,c2 = st.columns([2,3])
        ax.titulo = c1.text_input(f"Anexo {i+1} – Título", ax.titulo)
        ax.descricao = c2.text_input(f"Anexo {i+1} – Descrição", ax.descricao)
        ax.link = st.text_input(f"Anexo {i+1} – Link (opcional)", ax.link)

    rel.observacoes = st.text_area("Observações (opcional)", rel.observacoes)

    submitted = st.form_submit_button("Atualizar prévia")
    if submitted:
        st.session_state.rel = rel
        # Autosave local
        if st.session_state.get("autosave", True):
            try:
                p = Path(st.session_state.get("draft_dir", str(Path.cwd()/ "drafts")))
                p.mkdir(parents=True, exist_ok=True)
                name = f"{(rel.codigo or 'relatorio').replace(' ','_')}.json"
                (p / name).write_text(json.dumps(rel.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")
                st.toast("Rascunho salvo (local)", icon="💾")
            except Exception as e:
                st.warning(f"Autosave falhou: {e}")

        # Exportações
        md_bytes  = to_markdown(rel).encode("utf-8")
        pdf_bytes = build_pdf(rel, st.session_state.get("logo_bytes"), st.session_state.get("logo_width_cm", 3.5) if "logo_width_cm" in st.session_state else 3.5)
        docx_bytes= build_docx(rel, st.session_state.get("logo_bytes"), st.session_state.get("logo_width_cm", 3.5) if "logo_width_cm" in st.session_state else 3.5)

        # Uploads automáticos
        base = (rel.codigo or "relatorio").replace(" ", "_")
        if st.session_state.get("auto_drive", False):
            try:
                svc = get_drive_service()
                folder_id = st.secrets.get("drive", {}).get("folder_id")
                if not folder_id:
                    raise RuntimeError("Defina [drive].folder_id em st.secrets.")
                drive_upload_bytes(svc, folder_id, f"{base}.json", json.dumps(rel.model_dump(), ensure_ascii=False, indent=2).encode("utf-8"), "application/json")
                drive_upload_bytes(svc, folder_id, f"{base}.md", md_bytes, "text/markdown")
                drive_upload_bytes(svc, folder_id, f"{base}.pdf", pdf_bytes, "application/pdf")
                drive_upload_bytes(svc, folder_id, f"{base}.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.success("Upload automático → Google Drive concluído ✅")
            except Exception as e:
                st.error(f"Drive (auto): {e}")

        if st.session_state.get("auto_gh", False):
            try:
                import requests  # ensure available
                def gh_put(name, data, mime_hint):
                    url = f"https://api.github.com/repos/{st.secrets['github']['repo']}/contents/{st.secrets['github'].get('base_path','reports')}/{name}"
                    payload = {"message": f"auto: {name}", "content": base64.b64encode(data).decode("utf-8"), "branch": st.secrets['github'].get('branch','main')}
                    r = requests.put(url, json=payload, headers={"Authorization": f"token {st.secrets['github']['token']}", "Accept": "application/vnd.github+json"}, timeout=30)
                    if r.status_code not in (200,201): raise RuntimeError(r.text)
                gh_put(f"{base}.json", json.dumps(rel.model_dump(), ensure_ascii=False, indent=2).encode("utf-8"), "application/json")
                gh_put(f"{base}.md", md_bytes, "text/markdown")
                gh_put(f"{base}.pdf", pdf_bytes, "application/pdf")
                gh_put(f"{base}.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.success("Upload automático → GitHub concluído ✅")
            except Exception as e:
                st.error(f"GitHub (auto): {e}")

st.subheader("Prévia (Markdown)")
st.code(to_markdown(st.session_state.rel), language="markdown")

# Downloads manuais
colA, colB, colC = st.columns(3)
md_bytes = to_markdown(st.session_state.rel).encode("utf-8")
colA.download_button("⬇️ .md", md_bytes, file_name=f"{(st.session_state.rel.codigo or 'relatorio')}.md", mime="text/markdown", use_container_width=True)

try:
    pdf_bytes = build_pdf(st.session_state.rel, st.session_state.get("logo_bytes"), st.session_state.get("logo_width_cm", 3.5) if "logo_width_cm" in st.session_state else 3.5)
    colB.download_button("⬇️ PDF", pdf_bytes, file_name=f"{(st.session_state.rel.codigo or 'relatorio')}.pdf", mime="application/pdf", use_container_width=True)
except Exception as e:
    colB.error(f"PDF: {e}")

try:
    docx_bytes = build_docx(st.session_state.rel, st.session_state.get("logo_bytes"), st.session_state.get("logo_width_cm", 3.5) if "logo_width_cm" in st.session_state else 3.5)
    colC.download_button("⬇️ DOCX", docx_bytes, file_name=f"{(st.session_state.rel.codigo or 'relatorio')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
except Exception as e:
    colC.error(f"DOCX: {e}")
